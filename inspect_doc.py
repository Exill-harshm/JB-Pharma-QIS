import yaml
import docx
import re

def inspect():
    try:
        docx_path = "D:\\JB Pharma internal\\JB Pharma\\Cardiolek\\Output\\Cardiolek_QIS.docx"
        doc = docx.Document(docx_path)
        
        target_text_contains = "summary of controls performed at the critical steps"
        
        print("Searching for markers...")
        found_start_para = None
        for p in doc.paragraphs:
            if target_text_contains.lower() in p.text.lower():
                found_start_para = p
                print(f"START FOUND: '{p.text[:100]}'")
                break
        
        if not found_start_para:
            print("Start paragraph not found.")
            return

        # Find next paragraph starting with (b)
        found_end_para = None
        start_scanning = False
        for p in doc.paragraphs:
            if p == found_start_para:
                start_scanning = True
                continue
            if start_scanning:
                text = p.text.strip()
                if text.lower().startswith("(b)") or text.startswith("2.3."):
                    found_end_para = p
                    print(f"END FOUND: '{text[:100]}'")
                    break
        
        # Now find tables between these two
        # Use child elements to keep track of order
        elements = []
        for child in doc.element.body:
            if isinstance(child, docx.oxml.text.paragraph.CT_P):
                elements.append(docx.text.paragraph.Paragraph(child, doc))
            elif isinstance(child, docx.oxml.table.CT_Tbl):
                elements.append(docx.table.Table(child, doc))

        tables_found = []
        collecting = False
        for elem in elements:
            if isinstance(elem, docx.text.paragraph.Paragraph):
                if elem == found_start_para:
                    collecting = True
                    continue
                if elem == found_end_para:
                    break
            
            if collecting and isinstance(elem, docx.table.Table):
                tables_found.append(elem)

        print(f"Total tables found: {len(tables_found)}")
        for idx, tbl in enumerate(tables_found):
            print(f"\nTable {idx + 1}:")
            for r_idx, row in enumerate(tbl.rows[:2]):
                cells = [cell.text.strip() for cell in row.cells[:4]]
                print(f"Row {r_idx + 1}: {cells}")
                
    except Exception as e:
        print(f"Error: {e}")

inspect()
