from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from .models import ApiInfo, ManufactureInfo, P31ManufacturerInfo, SummaryInfo


class QisDocxFiller:
    @staticmethod
    def _row_cells(row):
        try:
            return tuple(row.cells)
        except Exception:
            return tuple()

    def fill(
        self,
        template_docx: Path,
        output_docx: Path,
        api_info: ApiInfo,
        summary_info: SummaryInfo | None = None,
        manufacture_info: ManufactureInfo | None = None,
        p31_info: P31ManufacturerInfo | None = None,
    ) -> list[str]:
        warnings: list[str] = []
        doc = Document(str(template_docx))

        if summary_info is not None:
            self._fill_summary_tables(doc, summary_info)

        if manufacture_info is not None:
            inserted = self._fill_manufacture_section(doc, manufacture_info)
            if not inserted:
                warnings.append("Could not locate 2.3.S.2.1 section in QIS template.")

        if p31_info is not None:
            inserted = self._fill_p31_section(doc, p31_info)
            if not inserted:
                warnings.append("Could not locate 2.3.P.3.1 section in QIS template.")

        table_found = self._fill_api_table(doc, api_info)
        if not table_found:
            inserted = self._insert_fallback_paragraphs(doc, api_info)
            if inserted:
                warnings.append("API table not found; inserted fallback paragraphs near 2.3.S heading.")
            else:
                warnings.append("API table and fallback anchor not found; output left unchanged.")

        destination = self._safe_output_path(output_docx)
        doc.save(str(destination))
        if destination != output_docx:
            warnings.append(f"Output file was locked; wrote to {destination}")

        return warnings

    def _fill_summary_tables(self, doc: Document, summary: SummaryInfo) -> None:
        if not doc.tables:
            return

        table = doc.tables[0]

        # Clear only rows that have extracted source values.
        # This preserves static merged rows such as section headings.
        for row_index, row in enumerate(table.rows):
            row_cells = self._row_cells(row)
            if not row_cells:
                continue
            label_text = row_cells[0].text
            values = self._lookup_summary_values(label_text, summary.summary_values_by_label)
            if not values:
                continue
            for col_index in range(1, len(row_cells)):
                row_cells[col_index].text = ""

        for row_index, row in enumerate(table.rows):
            row_cells = self._row_cells(row)
            if not row_cells:
                continue
            label_text = row_cells[0].text
            values = self._lookup_summary_values(label_text, summary.summary_values_by_label)
            if not values:
                continue

            for col_index in range(1, len(row_cells)):
                source_index = col_index - 1
                if source_index >= len(values):
                    continue
                self._set_text(table, row_index, col_index, values[source_index])

        self._apply_grouped_postal_address_placeholders(table, summary.summary_values_by_label)

        if len(doc.tables) > 1:
            self._fill_admin_summary_table(doc.tables[1], summary.summary_values_by_label)

        if len(doc.tables) > 2 and summary.related_row_values:
            related_table = doc.tables[2]
            if len(related_table.rows) > 1:
                second_row_cells = self._row_cells(related_table.rows[1])
                for col_index in range(len(second_row_cells)):
                    second_row_cells[col_index].text = ""
            for col_index, value in enumerate(summary.related_row_values[:4]):
                self._set_text(related_table, 1, col_index, value)
            if len(related_table.rows) > 2:
                row = related_table.rows[2]._tr
                row.getparent().remove(row)

    def _lookup_summary_values(self, target_label: str, source_map: dict[str, list[str]]) -> list[str]:
        target_key = self._normalize_label(target_label)
        if not target_key:
            return []

        direct = source_map.get(target_key)
        if direct:
            return direct

        best_key = ""
        best_score = 0
        for source_key, values in source_map.items():
            if not values:
                continue
            if source_key in target_key or target_key in source_key:
                score = min(len(source_key), len(target_key))
                if score > best_score:
                    best_key = source_key
                    best_score = score

        if best_key:
            return source_map[best_key]
        return []

    def _fill_admin_summary_table(self, table, source_map: dict[str, list[str]]) -> None:
        for row_index, row in enumerate(table.rows):
            row_cells = self._row_cells(row)
            if len(row_cells) < 2:
                continue
            label_text = row_cells[0].text
            values = self._lookup_summary_values(label_text, source_map)
            if not values:
                continue
            candidate = values[0].strip()
            if not candidate:
                continue
            row_cells[1].text = candidate

    def _apply_grouped_postal_address_placeholders(self, table, source_map: dict[str, list[str]]) -> None:
        self._apply_grouped_placeholder_block(
            table,
            source_map,
            [
                "Building/PO Box number",
                "Road/Street",
                "Plant/Zone",
                "Village/suburb",
            ],
        )
        self._apply_grouped_placeholder_block(
            table,
            source_map,
            [
                "Town/City",
                "District and Mandal",
                "Province/State",
                "Postal code",
            ],
        )

    def _apply_grouped_placeholder_block(
        self,
        table,
        source_map: dict[str, list[str]],
        label_group: list[str],
    ) -> None:
        if not label_group or len(table.rows) == 0:
            return

        label_to_row: dict[str, int] = {}
        for row_index, row in enumerate(table.rows):
            row_cells = self._row_cells(row)
            if not row_cells:
                continue
            row_label_key = self._normalize_label(row_cells[0].text)
            if row_label_key and row_label_key not in label_to_row:
                label_to_row[row_label_key] = row_index

        row_indices: list[int] = []
        values: list[str] = []
        for raw_label in label_group:
            key = self._normalize_label(raw_label)
            row_index = label_to_row.get(key)
            if row_index is None:
                return
            row_indices.append(row_index)

            source_values = source_map.get(key, [])
            if source_values:
                values.append(source_values[0].strip())
            else:
                values.append("")

        start_row = min(row_indices)
        end_row = max(row_indices)
        if start_row == end_row:
            return

        start_row_cells = self._row_cells(table.rows[start_row])
        if len(start_row_cells) < 2:
            return

        # Clear existing values in the grouped range before merging.
        for row_index in row_indices:
            row = table.rows[row_index]
            row_cells = self._row_cells(row)
            for col_index in range(1, len(row_cells)):
                row_cells[col_index].text = ""

        for col_index in range(1, len(start_row_cells)):
            try:
                table.cell(start_row, col_index).merge(table.cell(end_row, col_index))
            except Exception:
                pass

        grouped_value = "\n".join(values)
        table.cell(start_row, 1).text = grouped_value

    @staticmethod
    def _normalize_label(label: str) -> str:
        normalized = label.lower().replace("\n", " ")
        normalized = re.sub(r"(?<=\D)\d+\b", "", normalized)
        normalized = re.sub(r"[^a-z0-9]+", " ", normalized)
        return re.sub(r"\s+", " ", normalized).strip()

    @staticmethod
    def _set_text(table, row_index: int, col_index: int, value: str) -> None:
        if value is None:
            return
        if not str(value).strip():
            return
        if row_index >= len(table.rows):
            return
        row = table.rows[row_index]
        row_cells = QisDocxFiller._row_cells(row)
        if col_index >= len(row_cells):
            return
        row_cells[col_index].text = value

    def _fill_manufacture_section(self, doc: Document, info: ManufactureInfo) -> bool:
        heading_index = None
        for index, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text.startswith("2.3.S.2.1 Manufacturer(s)"):
                heading_index = index
                break

        # If legacy "Refer Section" placeholder exists, clear it.
        if heading_index is not None:
            for index, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if index > heading_index and text == "Refer Section 3.2.S.2.1":
                    paragraph.text = ""
                    break

        table = self._find_s2_manufacturer_table(doc)
        if table is None:
            return False

        self._set_text(table, 1, 0, info.name_and_address)
        self._set_text(table, 1, 1, info.responsibility)
        self._set_text(table, 1, 2, info.api_pq_number)
        self._set_text(table, 1, 3, info.letter_of_access)
        while len(table.rows) > 2:
            row = table.rows[2]._tr
            row.getparent().remove(row)

        return True

    @staticmethod
    def _find_s2_manufacturer_table(doc: Document):
        for table in doc.tables:
            if len(table.columns) != 4 or len(table.rows) < 2:
                continue
            header_cells = QisDocxFiller._row_cells(table.rows[0])
            header = " ".join(cell.text.strip().lower() for cell in header_cells)
            if "name and address" in header and "api-pq" in header and "letter of access" in header:
                return table
        return None

    def _fill_p31_section(self, doc: Document, info: P31ManufacturerInfo) -> bool:
        heading_index = None
        heading_paragraph = None
        refer_paragraph = None

        for index, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text == "2.3.P.3.1 Manufacturer(s)":
                heading_index = index
                heading_paragraph = paragraph
                continue
            if heading_index is not None and index > heading_index and text == "Refer Section 3.2.P.3.1":
                refer_paragraph = paragraph
                break

        if heading_paragraph is None:
            return False

        heading_paragraph.text = info.section_heading
        if refer_paragraph is not None:
            refer_paragraph.text = ""

        target_table = None
        for table in doc.tables:
            if len(table.columns) != 2 or len(table.rows) < 2:
                continue
            header_cells = self._row_cells(table.rows[0])
            if len(header_cells) < 2:
                continue
            header_left = header_cells[0].text.strip().lower()
            header_right = header_cells[1].text.strip().lower()
            if "name and address" in header_left and "responsibility" in header_right:
                target_table = table
                break

        if target_table is None:
            return False

        self._set_text(target_table, 1, 0, info.name_and_address)
        self._set_text(target_table, 1, 1, info.responsibility)
        while len(target_table.rows) > 2:
            row = target_table.rows[2]._tr
            row.getparent().remove(row)

        return True

    def _fill_api_table(self, doc: Document, api_info: ApiInfo) -> bool:
        found_api = False
        found_mfr = False
        full_details_row: tuple[int, int] | None = None

        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                row_cells = self._row_cells(row)
                if not row_cells:
                    continue
                row_text = " ".join(cell.text.strip() for cell in row_cells).lower()

                if "name of api" in row_text:
                    target_cell = row_cells[-1]
                    target_cell.text = api_info.api_name
                    found_api = True

                if "name of api manufacturer" in row_text:
                    target_cell = row_cells[-1]
                    target_cell.text = api_info.manufacturer_text
                    found_mfr = True

                if "full details in the pd" in row_text:
                    full_details_row = (ti, ri)

                if "confirmation of api prequalification document" in row_text and row_cells:
                    row_cells[0].text = "□"
                if "certificate of suitability to the european pharmacopoeia" in row_text and row_cells:
                    row_cells[0].text = "□"
                if "active pharmaceutical ingredient master file" in row_text and row_cells:
                    row_cells[0].text = "□"

        if full_details_row is not None:
            ti, ri = full_details_row
            chosen_row_cells = self._row_cells(doc.tables[ti].rows[ri])
            if chosen_row_cells:
                chosen_row_cells[0].text = "√"

        return found_api and found_mfr

    @staticmethod
    def _insert_fallback_paragraphs(doc: Document, api_info: ApiInfo) -> bool:
        marker = "2.3.s drug substance"
        for para in doc.paragraphs:
            if marker in para.text.lower():
                para.text = f"Name of API: {api_info.api_name}\nName of API manufacturer: {api_info.manufacturer_text}\nOption selected: Full details in the PD"
                return True
        return False

    @staticmethod
    def _safe_output_path(output_docx: Path) -> Path:
        try:
            output_docx.parent.mkdir(parents=True, exist_ok=True)
            with open(output_docx, "ab"):
                pass
            return output_docx
        except PermissionError:
            from datetime import datetime

            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            return output_docx.with_name(f"{output_docx.stem}_{stamp}{output_docx.suffix}")
