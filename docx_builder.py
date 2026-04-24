"""
Module: docx_builder
Responsibility: Opens QIS template DOCX, finds all "Refer Section X.X.X"
placeholders, injects extracted PDF content (text + tables),
cleans injected noise (headers/footers/page numbers), saves output DOCX.

Noise detection is fully automatic â€” no hardcoded company names.
The auto-detected blocklist comes from pdf_extractor._build_noise_blocklist().
"""
import re
import copy
import os
import docx
from docx.oxml import OxmlElement
import fitz
from docx.shared import RGBColor, Inches
from typing import Dict, Set
from logger_setup import get_logger

PLACEHOLDER_PATTERN = re.compile(
    r"Refer\s+(?:the\s+)?section\s+(\d+(?:\.[a-zA-Z0-9]+)+)",
    re.IGNORECASE
)
TEMPLATE_SECTION_PATTERN = re.compile(
    r"^\s*(2\.3\.[SP](?:\.[A-Za-z0-9]+)*)\b",
    re.IGNORECASE
)

MANUAL_ENTRY_SECTIONS = {'1.2', '1.3', '1.4', '1.5', '1.5.1', '1.5.2', '1.6'}
OVERLAY_MANAGED_SECTIONS = {'3.2.S.2.1', '3.2.P.3.1'}

# Generic regex patterns â€” work for ANY document, no company-specific strings
_PAGE_NUM_RE = re.compile(r'^\d{1,4}$')
_PAGE_OF_RE  = re.compile(r'^\d+\s+of\s+\d+\s*$', re.IGNORECASE)

def _remove_noise_tables(doc, logger) -> int:
    body = doc.element.body
    _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def _wt(elem):
        return ''.join(t.text or '' for t in elem.iter(f'{{{_NS}}}t')).strip()

    # All paragraph texts â€” used to detect 1x1 header tables
    all_para_texts = {p.text.strip() for p in doc.paragraphs if p.text.strip()}

    def _is_noise(elem):
        text = _wt(elem)
        clean = re.sub(r'\s+', '', text)
        rows = len(elem.findall(f'.//{{{_NS}}}tr'))
        cols = len(elem.findall(f'{{{_NS}}}tr/{{{_NS}}}tc'))
        lower_text = text.lower()

        # Never delete likely specification/result tables.
        protected_markers = (
            'specification',
            'acceptance criteria',
            'analytical procedure',
            'test',
            'impurities',
            'assay',
            'description',
            'identification',
        )
        if any(marker in lower_text for marker in protected_markers):
            return False

        if not clean:                                                   return True  # empty
        if re.match(r'^\d{1,4}$', text):                               return True  # "42"
        if re.match(r'^\d+\s+of\s+\d+$', text, re.I):                 return True  # "3 of 6"
        if len(text) < 90 and re.match(r'^.{5,75}\s+\d{1,4}$', text): return True  # "Co. 52"
        if rows == 1 and cols == 1 and len(text) < 100 and text in all_para_texts:
            return True  # 1x1 header table duplicated from a paragraph
        # DMF page-header tables ("Drug Mater File Version: 3.1...")
        if rows <= 4 and 'Drug Mater File' in text:
            return True
        # PDF section-header tables ("3.2.P PARTICULARS OF FINSHED...")
        if rows <= 3 and 'PARTICULARS' in text:
            return True
        return False

    to_remove = [e for e in list(body) if e.tag.split('}')[-1] == 'tbl' and _is_noise(e)]
    for elem in to_remove:
        body.remove(elem)
    if to_remove:
        logger.info(f"Removed {len(to_remove)} noise table(s).")
    return len(to_remove)


def _collapse_blank_paragraphs(doc, logger) -> int:
    """Collapses runs of consecutive blank paragraphs down to at most one."""
    body = doc.element.body
    _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    consecutive, to_remove = 0, []
    for elem in list(body):
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            text = ''.join(t.text or '' for t in elem.iter(f'{{{_NS}}}t')).strip()
            if not text:
                consecutive += 1
                if consecutive > 1:
                    to_remove.append(elem)
            else:
                consecutive = 0
        else:
            consecutive = 0
    for elem in to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
    if to_remove:
        logger.info(f"Removed {len(to_remove)} excess blank paragraph(s).")
    return len(to_remove)

def _fix_zero_width_tables(doc, logger) -> int:
    """Fixes tables injected by pdf2docx with tblW=0 â€” they render as blank boxes."""
    _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    PAGE_WIDTH_DXA = '9026'  # A4 with standard margins
    fixed = 0
    for elem in doc.element.body:
        if elem.tag.split('}')[-1] != 'tbl':
            continue
        tblPr = elem.find(f'{{{_NS}}}tblPr')
        if tblPr is None:
            continue
        tblW = tblPr.find(f'{{{_NS}}}tblW')
        if tblW is not None and tblW.get(f'{{{_NS}}}w', '') == '0':
            tblW.set(f'{{{_NS}}}w', PAGE_WIDTH_DXA)
            tblW.set(f'{{{_NS}}}type', 'dxa')
            fixed += 1
    if fixed:
        logger.info(f"Fixed {fixed} zero-width table(s).")
    return fixed

def _remove_repeated_header_paragraphs(doc, logger) -> int:
    """
    Removes PDF page-header paragraphs injected repeatedly by pdf2docx.
    Detection is fully automatic: any paragraph text that appears 3+ times
    AND is under 120 chars is treated as a repeating noise header.
    """
    from collections import Counter
    _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    freq = Counter(p.text.strip() for p in doc.paragraphs if p.text.strip())
    noise = {t for t, c in freq.items() if c >= 3 and len(t) < 120}

    removed = 0
    for elem in list(doc.element.body):
        if elem.tag.split('}')[-1] != 'p':
            continue
        text = ''.join(t.text or '' for t in elem.iter(f'{{{_NS}}}t')).strip()
        if text in noise:
            elem.getparent().remove(elem)
            removed += 1

    if removed:
        logger.info(f"Removed {removed} repeated header paragraph(s).")
    return removed

def _remove_pdf_noise_paragraphs(doc, logger) -> int:
    """Removes PDF-injected noise paragraphs not caught by frequency filter."""
    _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body
    elements = list(body)

    def _wt(e):
        return ''.join(t.text or '' for t in e.iter(f'{{{_NS}}}t')).strip()

    removed = 0
    for i, elem in enumerate(elements):
        if elem.tag.split('}')[-1] != 'p':
            continue
        text = _wt(elem)
        if not text:
            continue

        drop = False
        # 1. Company name + bare page number  e.g. "Starry Co. Ltd 230"
        if re.match(r'^.{10,75}\s+\d{1,4}$', text) and len(text) < 80:
            drop = True
        # 2. PDF section heading (3.2.X format) when template heading (2.3.X) exists nearby above
        elif re.search(r'3\.2[\. ][A-Z0-9P]', text) and len(text) < 200:
            for j in range(max(0, i - 10), i):
                if re.search(r'2\.3\.[SP]', _wt(elements[j])):
                    drop = True
                    break
        # 3. OCR artifacts and known single-occurrence PDF header lines
        elif text in {'FINIHSED PRODUCT SPECIFICATION Product name', 'C~CkedBY:'}:
            drop = True

        if drop:
            elem.getparent().remove(elem)
            removed += 1

    if removed:
        logger.info(f"Removed {removed} PDF noise paragraph(s).")
    return removed

def _remove_empty_visual_tables(doc, logger) -> int:
    """
    Removes tables that have:
    - no meaningful text
    - mostly empty cells
    - created from vector drawings (pdf2docx artifacts)
    """
    _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    removed = 0

    def get_text(cell):
        return ''.join(
            t.text or '' for t in cell._element.iter(f'{{{_NS}}}t')
        ).strip()

    for table in list(doc.tables):
        total_cells = 0
        empty_cells = 0
        text_cells = 0

        for row in table.rows:
            for cell in row.cells:
                total_cells += 1
                text = get_text(cell)

                if not text:
                    empty_cells += 1
                else:
                    text_cells += 1

        # KEY LOGIC
        if total_cells > 4 and text_cells == 0:
            # completely empty table â†’ remove
            table._element.getparent().remove(table._element)
            removed += 1

        elif total_cells > 6 and (empty_cells / total_cells) > 0.8:
            # mostly empty â†’ also remove
            table._element.getparent().remove(table._element)
            removed += 1

    if removed:
        logger.info(f"Removed {removed} empty visual table(s).")

    return removed


def _remove_low_content_injected_tables(doc, logger, keep_first_n_tables: int) -> int:
    """
    Removes low-value injected tables while preserving original template tables.
    This targets pdf2docx line-art table artifacts that appear as random lines.
    """
    _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    removed = 0

    def get_text(cell):
        return ''.join(
            t.text or '' for t in cell._element.iter(f'{{{_NS}}}t')
        ).strip()

    all_tables = list(doc.tables)
    for idx, table in enumerate(all_tables):
        if idx < keep_first_n_tables:
            continue

        total_cells = 0
        text_cells = 0
        text_chars = 0
        for row in table.rows:
            for cell in row.cells:
                total_cells += 1
                text = get_text(cell)
                if text:
                    text_cells += 1
                    text_chars += len(text)

        if total_cells == 0:
            continue

        empty_ratio = (total_cells - text_cells) / total_cells
        should_remove = False
        if text_cells == 0 and total_cells >= 4:
            should_remove = True
        elif empty_ratio > 0.85 and text_chars < 80 and total_cells >= 6:
            should_remove = True
        elif text_cells <= 2 and total_cells >= 8 and text_chars < 100:
            should_remove = True

        if should_remove:
            table._element.getparent().remove(table._element)
            removed += 1

    if removed:
        logger.info(f"Removed {removed} low-content injected table artifact(s).")
    return removed


def _is_noise_paragraph(text: str, blocklist: Set[str]) -> bool:
    """
    Returns True if text is a header/footer/page-number noise line.

    Uses two layers:
    1. Auto-detected blocklist (strings that repeat in top/bottom margins)
    2. Generic document-agnostic heuristics (page numbers, short ALL-CAPS)
    """
    text = text.strip()
    if not text:
        return False

    norm = " ".join(text.lower().split())

    # Layer 1: auto-detected repeating header/footer text
    if norm in blocklist:
        return True

    # Layer 2: bare page numbers  ("42"  or  "42 of 100")
    if _PAGE_NUM_RE.match(norm):
        return True
    if _PAGE_OF_RE.match(norm):
        return True

    # Layer 2b: OCR/header garbage lines seen in converted PDFs.
    compact = re.sub(r'[^a-z0-9]+', '', norm)
    if 'productspecificationproductname' in compact:
        return True
    if compact in {'cckedby', 'checkedby'}:
        return True

    # Layer 3: very short ALL-CAPS lines â€” running header artefacts
    # (e.g. "INTRODUCTION", "MODULE 3") â€” only strip if <= 3 words & < 25 chars
    if len(text) < 25 and text.isupper() and len(text.split()) <= 3:
        return True

    return False


def _is_footer_table_row(row, blocklist: Set[str]) -> bool:
    """
    Returns True if a table row is entirely noise (header/footer row).

    A footer row must satisfy:
    - At least one cell matches the auto-detected blocklist, AND
    - Every other non-empty cell is a bare page number.

    This is fully generic â€” no company names needed.
    """
    cell_texts = [cell.text.strip() for cell in row.cells]
    if not any(cell_texts):
        return False

    has_blocklist_hit = False
    for t in cell_texts:
        if not t:
            continue
        norm = " ".join(t.lower().split())
        if norm in blocklist:
            has_blocklist_hit = True
        elif not (_PAGE_NUM_RE.match(t) and t.isdigit() and int(t) < 10000):
            # This cell has real content â€” row is NOT a footer
            return False

    return has_blocklist_hit


def _clean_injected_content(
    src_doc,
    blocklist: Set[str],
    logger,
    section_num: str
):
    """
    Cleans noise from pdf2docx-converted DOCX BEFORE injecting into template.
    Uses the auto-detected blocklist â€” no hardcoded company strings.
    """
    removed = 0

    # Remove noisy standalone paragraphs entirely to avoid leaving empty lines.
    paras_to_remove = []
    for para in src_doc.paragraphs:
        if _is_noise_paragraph(para.text, blocklist):
            paras_to_remove.append(para)

    for para in paras_to_remove:
        p = para._p
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)
            removed += 1

    for table in src_doc.tables:
        for row in table.rows:
            if _is_footer_table_row(row, blocklist):
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.clear()
                removed += 1

    if removed:
        logger.info(
            f"Section {section_num}: cleaned {removed} "
            f"noise items (headers/footers/page numbers)."
        )

    if section_num == "3.2.S.6":
        _trim_s6_container_closure_content(src_doc, logger)
    elif section_num == "3.2.P.3.3":
        _trim_p33_narrative_content(src_doc, logger)


def _trim_s6_container_closure_content(src_doc, logger) -> None:
    """
    For 3.2.S.6, keep only one concise descriptive paragraph in the template.
    This prevents spillover of 3.2.S.6.1/3.2.S.6.2 and avoids carrying 3.2-style
    headings into the 2.3 summary narrative.
    """
    body = src_doc.element.body

    def _is_candidate(text: str) -> bool:
        text = " ".join(text.split())
        if not text:
            return False
        low = text.lower()

        if re.match(r'^\d+(\s+of\s+\d+)?$', low):
            return False
        if low.startswith("3.2.s.6"):
            return False
        if any(tok in low for tok in ("drug mater file", "version:", "module:")):
            return False
        if re.search(r"\bfigure\b", low):
            return False
        if re.search(r"\btable\b", low):
            return False

        # Keep the first substantial descriptive sentence block.
        return len(text) >= 90 and text.count(" ") >= 12

    keep_p = None
    for para in src_doc.paragraphs:
        if _is_candidate(para.text):
            keep_p = para._p
            break

    if keep_p is None:
        return

    removed = 0
    for elem in list(body):
        tag_name = elem.tag.split('}')[-1]
        if tag_name == 'p' and elem == keep_p:
            continue
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
            removed += 1

    if removed:
        logger.info(f"Section 3.2.S.6: trimmed injected content to a single summary paragraph.")


def _trim_p33_narrative_content(src_doc, logger) -> None:
    """
    For 3.2.P.3.3 keep narrative text only and remove flow-chart tail labels.
    The flow diagram itself is inserted separately as a rendered image.
    """
    body = src_doc.element.body
    cut_started = False
    removed = 0

    for elem in list(body):
        if elem.tag.split('}')[-1] != 'p':
            continue

        text = " ".join(_element_text_content(elem).split())
        low = text.lower()

        if low.startswith("3.2.p.3.3 description of manufacturing process"):
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
                removed += 1
            continue

        if re.search(r"\b(manufacturing process flow chart|flow\s+diagram)\b", low):
            cut_started = True

        if cut_started:
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
                removed += 1

    if removed:
        logger.info("Section 3.2.P.3.3: trimmed flow-chart tail text; image will be inserted separately.")


def _extract_p33_flow_diagram_image(pdf_path: str, log_folder: str, logger) -> str:
    """Renders the 3.2.P.3.3 flow diagram area from source PDF into a PNG."""
    doc = None
    try:
        doc = fitz.open(pdf_path)
        target_page = None
        for i in range(doc.page_count - 1, -1, -1):
            page = doc[i]
            text = page.get_text("text", sort=True).lower()
            if (
                "manufacturing process flow chart" in text
                or ("dispensing" in text and "packing" in text and "filtration" in text)
            ):
                target_page = page
                break

        if target_page is None:
            logger.warning("Section 3.2.P.3.3: could not locate flow-chart page in source PDF.")
            return ""

        page = target_page
        page_rect = page.rect

        top_hits = page.search_for("MANUFACTURING PROCESS FLOW CHART")
        if not top_hits:
            top_hits = page.search_for("Flow diagram of the manufacturing process")
        top_y = top_hits[0].y0 - 8 if top_hits else page_rect.y0 + page_rect.height * 0.08

        bottom_y = None
        for token in ("Packing", "Labeling", "Visual Inspection", "Terminal Sterilization"):
            hits = page.search_for(token)
            if hits:
                y = max(r.y1 for r in hits) + 18
                bottom_y = y if bottom_y is None else max(bottom_y, y)

        if bottom_y is None:
            bottom_y = page_rect.y0 + page_rect.height * 0.92

        x0 = page_rect.x0 + page_rect.width * 0.04
        x1 = page_rect.x1 - page_rect.width * 0.04
        y0 = max(page_rect.y0, top_y)
        y1 = min(page_rect.y1, bottom_y)

        if y1 - y0 < 40:
            logger.warning("Section 3.2.P.3.3: flow-chart clip too small; using full-page body area.")
            y0 = page_rect.y0 + page_rect.height * 0.08
            y1 = page_rect.y0 + page_rect.height * 0.95

        clip = fitz.Rect(x0, y0, x1, y1)
        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), clip=clip, alpha=False)

        img_path = os.path.join(log_folder, "temp_3.2.P.3.3_flow_diagram.png")
        pix.save(img_path)
        logger.info(f"Section 3.2.P.3.3: extracted flow diagram image -> {img_path}")
        return img_path
    except Exception as e:
        logger.warning(f"Section 3.2.P.3.3: flow diagram image extraction failed: {e}")
        return ""
    finally:
        if doc is not None:
            doc.close()


def _insert_p33_flow_diagram_image(doc, image_path: str, logger) -> bool:
    """Inserts extracted flow diagram under the template flow-diagram prompt."""
    try:
        for para in doc.paragraphs:
            low = para.text.strip().lower()
            if "flow diagram of the manufacturing process" in low:
                run = para.add_run()
                run.add_break()
                run.add_picture(image_path, width=Inches(5.8))
                logger.info("Section 3.2.P.3.3: inserted flow diagram image under template prompt.")
                return True
    except Exception as e:
        logger.warning(f"Section 3.2.P.3.3: failed to insert flow diagram image: {e}")
    return False


def _element_text_content(element) -> str:
    """Extracts combined text for an XML element."""
    _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    return ''.join(t.text or '' for t in element.iter(f'{{{_NS}}}t')).strip()


def _iter_all_paragraphs(doc):
    """Yields every paragraph in body AND inside every table cell."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _insert_warning(paragraph, section: str):
    """Replaces placeholder with visible red WARNING text."""
    paragraph.clear()
    run = paragraph.add_run(
        f"[WARNING: Could not populate section {section} - check source PDF]"
    )
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)


def _strip_drawing_elements(element):
    """
    Removes drawing/image XML nodes using Clark notation.
    Images are re-inserted separately via PyMuPDF bytes.
    """
    DRAWING_TAGS = {
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing',
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict',
        '{urn:schemas-microsoft-com:vml}imagedata',
        '{urn:schemas-microsoft-com:vml}shape',
        '{urn:schemas-microsoft-com:office:office}OLEObject',
        '{http://schemas.openxmlformats.org/drawingml/2006/main}blipFill',
    }
    to_remove = [n for n in element.iter() if n.tag in DRAWING_TAGS]
    for node in to_remove:
        parent = node.getparent()
        if parent is not None:
            try:
                parent.remove(node)
            except Exception:
                pass

def _analyze_injected_doc_layout(src_doc) -> dict:
    """Computes simple structure metrics used to decide table auto-inclusion."""
    body_elements = [
        elem for elem in src_doc.element.body
        if elem.tag.split('}')[-1] != 'sectPr'
    ]

    table_count = 0
    rich_table_count = 0
    table_text_chars = 0
    nonblank_paragraph_count = 0

    for elem in body_elements:
        tag_name = elem.tag.split('}')[-1]
        if tag_name == 'tbl':
            table_count += 1
            tbl_text = _element_text_content(elem)
            table_text_chars += len(tbl_text)

            rows = len(elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'))
            first_row = elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')
            cols = len(first_row.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')) if first_row is not None else 0

            if rows >= 8 and cols >= 3 and len(tbl_text) >= 250:
                rich_table_count += 1
        elif tag_name == 'p':
            if _element_text_content(elem):
                nonblank_paragraph_count += 1

    return {
        "table_count": table_count,
        "rich_table_count": rich_table_count,
        "table_text_chars": table_text_chars,
        "nonblank_paragraph_count": nonblank_paragraph_count,
    }


def _should_auto_include_tables(layout: dict) -> bool:
    """
    Enables table injection for table-dominant converted sections.
    Keeps behavior generic and avoids section-specific hardcoding.
    """
    return (
        layout.get("rich_table_count", 0) >= 2
        and layout.get("table_count", 0) >= 2
        and layout.get("table_text_chars", 0) >= 1000
        and layout.get("nonblank_paragraph_count", 0) <= 4
    )


def _table_header_signature(table_element) -> str:
    """Builds a normalized signature from the first row of a table element."""
    _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    first_row = table_element.find(f'{{{_NS}}}tr')
    if first_row is None:
        return ""

    cells = first_row.findall(f'{{{_NS}}}tc')
    if not cells:
        return ""

    parts: list[str] = []
    for cell in cells:
        text = ''.join(t.text or '' for t in cell.iter(f'{{{_NS}}}t'))
        norm = re.sub(r'\s+', ' ', text).strip().lower()
        parts.append(norm)

    non_empty = [p for p in parts if p]
    if len(non_empty) < 2:
        return ""
    return '|'.join(parts)


def _table_column_count(table_element) -> int:
    _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    first_row = table_element.find(f'{{{_NS}}}tr')
    if first_row is None:
        return 0
    return len(first_row.findall(f'{{{_NS}}}tc'))


def _drop_consecutive_duplicate_table_headers(src_doc, logger, section_num: str) -> int:
    """
    Removes first-row headers that repeat across consecutive tables.
    This handles PDF page-split continuation tables cleanly.
    """
    _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    removed = 0
    last_header_signature = ""

    for element in src_doc.element.body:
        if element.tag.split('}')[-1] != 'tbl':
            continue

        signature = _table_header_signature(element)
        if not signature:
            continue

        rows = element.findall(f'{{{_NS}}}tr')
        if (
            last_header_signature
            and signature == last_header_signature
            and len(rows) > 1
        ):
            first_row = rows[0]
            parent = first_row.getparent()
            if parent is not None:
                parent.remove(first_row)
                removed += 1
            continue

        last_header_signature = signature

    if removed:
        logger.info(
            f"Section {section_num}: removed {removed} repeated continuation table header row(s)."
        )
    return removed


def _merge_consecutive_continuation_tables(src_doc, logger, section_num: str) -> int:
    """
    Merges consecutive tables that share the same column schema and are not
    separated by non-empty paragraphs. This smooths page-break table splits.
    """
    _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    merged = 0

    previous_table = None
    previous_cols = 0
    paragraph_between = False

    for elem in list(src_doc.element.body):
        tag_name = elem.tag.split('}')[-1]

        if tag_name == 'p':
            if _element_text_content(elem):
                paragraph_between = True
            continue

        if tag_name != 'tbl':
            previous_table = None
            previous_cols = 0
            paragraph_between = False
            continue

        cols = _table_column_count(elem)
        can_merge = (
            previous_table is not None
            and not paragraph_between
            and cols > 0
            and cols == previous_cols
        )

        if can_merge:
            for row in elem.findall(f'{{{_NS}}}tr'):
                previous_table.append(copy.deepcopy(row))

            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
            merged += 1
            paragraph_between = False
            continue

        previous_table = elem
        previous_cols = cols
        paragraph_between = False

    if merged:
        logger.info(
            f"Section {section_num}: merged {merged} continuation table(s) across page breaks."
        )
    return merged


def _drop_outlier_table_schemas(src_doc, logger, section_num: str) -> int:
    """
    Removes minority table schemas when one column layout clearly dominates.
    This keeps page-split continuation tables and drops unrelated outlier tables.
    """
    from collections import Counter

    table_elements = [
        elem for elem in src_doc.element.body
        if elem.tag.split('}')[-1] == 'tbl'
    ]
    if len(table_elements) < 3:
        return 0

    schema_by_table = []
    for elem in table_elements:
        first_row = elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')
        if first_row is None:
            continue
        cols = len(first_row.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'))
        if cols <= 0:
            continue
        schema_by_table.append((elem, cols))

    if len(schema_by_table) < 3:
        return 0

    freq = Counter(cols for _, cols in schema_by_table)
    dominant_cols, dominant_count = max(freq.items(), key=lambda item: item[1])

    if dominant_count < 2 or dominant_count == len(schema_by_table):
        return 0

    removed = 0
    for elem, cols in schema_by_table:
        if cols == dominant_cols:
            continue
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
            removed += 1

    if removed:
        logger.info(
            f"Section {section_num}: removed {removed} outlier table schema(s) "
            f"(dominant columns={dominant_cols})."
        )
    return removed


def _merge_split_tables(src_doc, logger, section_num: str) -> int:
    """
    Merge consecutive table fragments (typically split by PDF page breaks)
    when both parts have the same column count.
    """
    if section_num == "3.2.S.4.1":
        return 0

    body = src_doc.element.body
    elements = list(body)
    merged = 0

    def _tag_name(el) -> str:
        return el.tag.split('}')[-1]

    def _is_blank_paragraph(el) -> bool:
        if _tag_name(el) != 'p':
            return False
        _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        txt = ''.join(t.text or '' for t in el.iter(f'{{{_NS}}}t')).strip()
        return txt == ""

    i = 0
    while i < len(elements) - 1:
        first = elements[i]
        if _tag_name(first) != 'tbl':
            i += 1
            continue

        j = i + 1
        while j < len(elements) and _is_blank_paragraph(elements[j]):
            j += 1
        if j >= len(elements):
            break

        second = elements[j]
        if _tag_name(second) != 'tbl':
            i = j
            continue

        rows_a = first.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')
        rows_b = second.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')
        if not rows_a or not rows_b:
            i = j
            continue

        cols_a = len(rows_a[0].findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'))
        cols_b = len(rows_b[0].findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'))
        if cols_a != cols_b or cols_a == 0:
            i = j
            continue

        # If the second table starts with a header row, skip that row.
        def _row_text(row_el) -> str:
            _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            return " ".join(
                ''.join(t.text or '' for t in tc.iter(f'{{{_NS}}}t')).strip().lower()
                for tc in row_el.findall(f'{{{_NS}}}tc')
            )

        start_idx = 0
        first_row_text = _row_text(rows_b[0])
        if (
            'test' in first_row_text and
            ('acceptance' in first_row_text or 'criteria' in first_row_text) and
            ('method' in first_row_text or 'analytical procedure' in first_row_text)
        ):
            start_idx = 1

        for row_el in rows_b[start_idx:]:
            first.append(copy.deepcopy(row_el))

        # Remove merged second table and any blank separator paragraphs.
        body.remove(second)
        for k in range(i + 1, j):
            if k < len(elements) and _is_blank_paragraph(elements[k]):
                try:
                    body.remove(elements[k])
                except Exception:
                    pass

        merged += 1
        elements = list(body)
        i = max(i - 1, 0)

    if merged:
        logger.info(f"Section {section_num}: merged {merged} split table fragment(s).")
    return merged


def _clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\r", "\n")
    text = text.replace("\n", " ")
    text = re.sub(r"\s+", " ", text)
    text = text.strip()
    # Remove leading punctuation artifacts like ": value"
    text = re.sub(r'^\s*[:;\-]\s*', '', text)
    # Ensure a space after period when OCR glues words: "sample.The"
    text = re.sub(r'(?<=\w)\.(?=[A-Za-z])', '. ', text)
    # Start bullets/numbered points on a new line.
    text = re.sub(r'\s*[â€¢â—â–ªâ—¦]\s*', '\nâ€¢ ', text)
    text = re.sub(r'(?<!^)\s+(?=(\d{1,2}[.)]\s+))', '\n', text)
    text = re.sub(r'(?<!^)\s+(?=([A-Za-z][.)]\s+))', '\n', text)
    text = re.sub(r'(?<!^)\s*-\s+(?=[A-Za-z])', '\n- ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r"\s+", " ", text)
    # Keep intentional line breaks created above.
    text = re.sub(r' *\n *', '\n', text)
    return text.strip()


def _extract_scanned_spec_rows_with_camelot(pdf_path: str, logger):
    """
    Extracts scanned rows for section 2.3.S.4.1 using Camelot.
    Looks for pages near 'Lists of Tests and/& Specification'.
    Returns rows shaped as [Test, Acceptance criteria, Method].
    """
    try:
        import fitz
        import camelot
    except Exception as e:
        logger.warning(f"Camelot path unavailable for scanned table extraction: {e}")
        return []

    marker_patterns = [
        re.compile(r'lists?\s+of\s+tests?\s*(?:and|&|/)?\s*specifications?', re.IGNORECASE),
        re.compile(r'tests?\s*(?:&|and|/)\s*specifications?', re.IGNORECASE),
    ]

    marker_pages = []
    try:
        pdf = fitz.open(pdf_path)
        for i, page in enumerate(pdf):
            txt = " ".join(page.get_text().split())
            if any(p.search(txt) for p in marker_patterns):
                marker_pages.append(i + 1)  # Camelot uses 1-based pages
        pdf.close()
    except Exception as e:
        logger.warning(f"Could not scan marker pages in {os.path.basename(pdf_path)}: {e}")
        return []

    if not marker_pages:
        # Scanned PDFs often don't expose text to fitz; keep deterministic fallback.
        pages = "2,3,4"
        logger.info(
            f"{os.path.basename(pdf_path)}: marker not found in text layer; "
            f"using Camelot fallback pages {pages}."
        )
    else:
        page_set = set()
        for p in marker_pages:
            page_set.add(p)
            page_set.add(p + 1)
            page_set.add(p + 2)
        pages = ",".join(str(p) for p in sorted(page_set))

    try:
        tables = camelot.read_pdf(pdf_path, pages=pages, flavor="lattice")
    except Exception as e:
        logger.warning(f"Camelot failed on {os.path.basename(pdf_path)} pages {pages}: {e}")
        return []

    rows = []
    seen = set()
    seen_test_names = set()
    exclude_tokens = [
        "prepared by", "checked by", "approved by",
        "product name", "module", "ref. no", "drug mater file",
        "version", "date"
    ]

    for table in tables:
        df = table.df
        joined = " ".join(df.astype(str).values.flatten()).lower()
        if any(tok in joined for tok in exclude_tokens):
            continue

        table_rows = []
        for i in range(len(df)):
            raw = [str(x).strip() for x in df.iloc[i].tolist()]
            if all(not x for x in raw):
                continue

            if len(raw) >= 3:
                test = _clean_text(raw[1])
                spec = _clean_text(" ".join(raw[2:]))
            elif len(raw) == 2:
                test = _clean_text(raw[0])
                spec = _clean_text(raw[1])
            else:
                continue

            if test.lower() in {"", "tests", "test", "sr.", "no.", "s. no.", "s.no"}:
                continue
            if not spec:
                continue

            test_norm = re.sub(r'[^a-z0-9]+', '', test.lower())
            if test_norm in seen_test_names:
                continue

            rec = (test, spec, "")
            if rec not in seen:
                table_rows.append([test, spec, ""])
                seen_test_names.add(test_norm)

        if not table_rows:
            continue

        for row in table_rows:
            rec = (row[0], row[1], row[2])
            if rec not in seen:
                rows.append(row)
                seen.add(rec)

    if rows:
        logger.info(
            f"{os.path.basename(pdf_path)}: Camelot extracted {len(rows)} scanned rows."
        )
    return rows


def _append_rows_as_table(doc, current_anchor, headers, rows, force_new_table=False, insert_page_break=False):
    """
    Appends rows into existing anchor table when possible to preserve format.
    Falls back to creating a new table only if no suitable anchor table exists.
    Returns new anchor.
    """
    if not rows:
        return current_anchor

    # Find the preceding table for deduplication regardless of if we're appending to it.
    reference_table = None
    if hasattr(current_anchor, "tag") and current_anchor.tag.endswith("}tbl"):
        for t in doc.tables:
            if t._tbl == current_anchor and len(t.columns) >= len(headers):
                reference_table = t
                break

    anchor_table = reference_table if not force_new_table else None

    def _norm_key(s: str) -> str:
        return re.sub(r'[^a-z0-9]+', ' ', (s or '').lower()).strip()

    # Drop duplicates already present in reference table by first-column key.
    filtered_rows = rows[:]
    if reference_table is not None:
        existing_first_col = set()
        for r_idx, r in enumerate(reference_table.rows):
            if r_idx == 0:
                continue
            if not r.cells:
                continue
            k = _norm_key(r.cells[0].text)
            if k:
                existing_first_col.add(k)

        unique_rows = []
        seen_new = set()
        for row in filtered_rows:
            first_key = _norm_key(str(row[0]) if row else "")
            if not first_key:
                continue
            if first_key in existing_first_col or first_key in seen_new:
                continue
            unique_rows.append(row)
            seen_new.add(first_key)
        filtered_rows = unique_rows

    if not filtered_rows:
        return current_anchor

    if anchor_table is not None:
        for row in filtered_rows:
            cells = anchor_table.add_row().cells
            for j, val in enumerate(row):
                if j < len(cells):
                    cells[j].text = _clean_text(str(val))
        return current_anchor

    temp_table = doc.add_table(rows=len(filtered_rows) + 1, cols=len(headers))
    temp_table.style = "Table Grid"
    for j, col in enumerate(headers):
        run = temp_table.cell(0, j).paragraphs[0].add_run(col)
        run.bold = True
    for i, row in enumerate(filtered_rows, start=1):
        for j, val in enumerate(row):
            temp_table.cell(i, j).text = _clean_text(str(val))

    new_tbl = copy.deepcopy(temp_table._tbl)
    # Keep visual separation from previous content.
    gap_anchor = current_anchor
    for _ in range(4):
        new_p = OxmlElement('w:p')
        gap_anchor.addnext(new_p)
        gap_anchor = new_p
    gap_anchor.addnext(new_tbl)
    temp_table._tbl.getparent().remove(temp_table._tbl)
    return new_tbl


def _remove_tables_until_next_section(anchor_xml, logger, section_num: str) -> int:
    """
    Removes table scaffolds after current anchor until next section heading
    paragraph is reached (e.g., next 2.3.S/P.* heading).
    """
    removed = 0
    removed_paras = 0
    nxt = anchor_xml.getnext()
    _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    removed_other = 0

    while nxt is not None:
        tag = nxt.tag.split('}')[-1]
        if tag == 'p':
            text = ''.join(t.text or '' for t in nxt.iter(f'{{{_NS}}}t')).strip()
            if text and TEMPLATE_SECTION_PATTERN.search(text):
                break
            parent = nxt.getparent()
            to_remove = nxt
            nxt = nxt.getnext()
            if parent is not None:
                parent.remove(to_remove)
                removed_paras += 1
            continue
        elif tag == 'tbl':
            parent = nxt.getparent()
            to_remove = nxt
            nxt = nxt.getnext()
            if parent is not None:
                parent.remove(to_remove)
                removed += 1
            continue
        else:
            # Remove any other block-level leftovers in this section region.
            parent = nxt.getparent()
            to_remove = nxt
            nxt = nxt.getnext()
            if parent is not None:
                parent.remove(to_remove)
                removed_other += 1
            continue
        nxt = nxt.getnext()

    if removed or removed_paras or removed_other:
        logger.info(
            f"Section {section_num}: removed {removed} scaffold table(s) and "
            f"{removed_paras} scaffold paragraph(s) and "
            f"{removed_other} other scaffold block(s)."
        )
    return removed


def _add_section_spacing(anchor_xml, lines: int = 2):
    """Adds blank paragraph lines after current section content."""
    curr = anchor_xml
    for _ in range(max(0, lines)):
        p = OxmlElement('w:p')
        curr.addnext(p)
        curr = p
    return curr


def _inject_docx_content(
    src_docx_path: str,
    anchor_p_xml,
    blocklist: Set[str],
    logger,
    section_num: str,
    include_pdf_tables: bool = False,
    table_only: bool = False,
    table_text_keyword: str = ""
):
    """
    Opens pdf2docx temp DOCX, cleans noise, copies body elements
    into template after anchor. Skips sectPr and strips drawing refs.
    Returns new anchor (last inserted element).
    """
    current_anchor = anchor_p_xml
    try:
        src_doc = docx.Document(src_docx_path)
    except Exception as e:
        logger.error(
            f"Cannot open converted DOCX for section {section_num}: {e}"
        )
        return current_anchor

    _clean_injected_content(src_doc, blocklist, logger, section_num)
    _merge_split_tables(src_doc, logger, section_num)

    effective_include_pdf_tables = include_pdf_tables
    suppress_paragraphs = False
    if not include_pdf_tables:
        layout = _analyze_injected_doc_layout(src_doc)
        if _should_auto_include_tables(layout):
            effective_include_pdf_tables = True
            suppress_paragraphs = True
            _drop_outlier_table_schemas(src_doc, logger, section_num)
            _drop_consecutive_duplicate_table_headers(src_doc, logger, section_num)
            _merge_consecutive_continuation_tables(src_doc, logger, section_num)
            logger.info(
                f"Section {section_num}: detected table-dominant content "
                f"(tables={layout['table_count']}, rich_tables={layout['rich_table_count']}, "
                f"table_chars={layout['table_text_chars']}). Injecting tables automatically."
            )

    inserted_nonblank = False
    pending_blank_para = None

    table_count = 0
    keyword = table_text_keyword.strip().lower()

    def _element_text(element) -> str:
        _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        return ''.join(t.text or '' for t in element.iter(f'{{{_NS}}}t')).strip()

    body_elements = list(src_doc.element.body)

    def _prev_non_empty_paragraph_text(elements, idx: int) -> str:
        _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        for k in range(idx - 1, -1, -1):
            el = elements[k]
            tag = el.tag.split('}')[-1]
            if tag != 'p':
                continue
            txt = ''.join(t.text or '' for t in el.iter(f'{{{_NS}}}t')).strip()
            if txt:
                return txt
        return ""

    for idx, element in enumerate(body_elements):
        if element.tag.endswith('}sectPr') or element.tag == 'sectPr':
            continue

        is_table = element.tag.endswith('}tbl') or element.tag == 'tbl'
        tag_name = element.tag.split('}')[-1]

        if table_only:
            if not is_table:
                continue
        elif not effective_include_pdf_tables and is_table:
            continue

        if table_only and keyword and is_table:
            table_text = " ".join(_element_text(element).split())
            table_text_lower = table_text.lower()
            caption_text = _prev_non_empty_paragraph_text(body_elements, idx)
            caption_text_lower = " ".join(caption_text.lower().split())

            # Exclude DMF header metadata tables.
            is_dmf_header_table = (
                'drug mater file' in table_text_lower or
                (
                    'product name' in table_text_lower and
                    'module:' in table_text_lower and
                    'version' in table_text_lower
                )
            )
            if is_dmf_header_table:
                continue

            has_keyword = (
                (keyword in table_text_lower) or
                (keyword in caption_text_lower)
            )
            if not has_keyword:
                continue

        if tag_name == 'p':
            if suppress_paragraphs:
                continue
            if not _element_text_content(element):
                # Drop leading blanks and compress internal blank runs.
                if not inserted_nonblank:
                    continue
                pending_blank_para = copy.deepcopy(element)
                continue

            # Insert at most one blank paragraph before non-blank content.
            if pending_blank_para is not None:
                try:
                    blank_el = copy.deepcopy(pending_blank_para)
                    current_anchor.addnext(blank_el)
                    current_anchor = blank_el
                except Exception:
                    pass
                pending_blank_para = None
        try:
            new_el = copy.deepcopy(element)
            _strip_drawing_elements(new_el)
            current_anchor.addnext(new_el)
            current_anchor = new_el
            if tag_name == 'p' and _element_text_content(new_el):
                inserted_nonblank = True
            if is_table:
                table_count += 1
        except Exception as el_e:
            logger.warning(
                f"Skipped element for section {section_num}: {el_e}"
            )

    if table_only and table_count == 0:
        logger.warning(
            f"Section {section_num}: no table containing keyword "
            f"'{table_text_keyword}' found. "
            f"No fallback to full content."
        )

    return current_anchor


def process_template(
    template_path:       str,
    output_path:         str,
    section_map:         Dict[str, str],
    log_folder:          str,
    section_page_limits: Dict[str, int] = None,
    section_start_pages: Dict[str, int] = None,
    preserve_template_tables: bool = False,
    include_pdf_tables: bool = False,
    table_only_sections: Set[str] = None,
    table_only_all_sections: bool = False,
    table_keyword_by_template_section: Dict[str, str] = None,
):
    """
    Main pipeline:
    1. Open QIS template DOCX
    2. Scan all paragraphs + table cells for placeholders
    3. Extract + clean + inject content per section
    4. Save populated output DOCX
    Returns: (sections_filled, warnings_count, failures)
    """
    logger = get_logger(log_folder)
    from pdf_extractor import extract_pdf_content

    if section_page_limits is None:
        section_page_limits = {}
    if section_start_pages is None:
        section_start_pages = {}
    if table_only_sections is None:
        table_only_sections = set()
    if table_keyword_by_template_section is None:
        table_keyword_by_template_section = {}

    try:
        doc = docx.Document(template_path)
    except Exception as e:
        logger.error(f"Failed to open template: {e}")
        raise

    template_table_count = len(doc.tables)

    sections_filled    = 0
    warnings_count     = 0
    failures           = 0
    processed_sections = set()

    logger.info("Starting QIS template placeholder scan.")
    body_paragraphs = list(doc.paragraphs)

    def _find_template_section_for_paragraph(target_para) -> str:
        idx = -1
        for i, p in enumerate(body_paragraphs):
            if p._p == target_para._p:
                idx = i
                break
        if idx == -1:
            return ""
        for j in range(idx, -1, -1):
            m = TEMPLATE_SECTION_PATTERN.search(body_paragraphs[j].text or "")
            if m:
                return m.group(1)
        return ""

    for paragraph in _iter_all_paragraphs(doc):
        match = PLACEHOLDER_PATTERN.search(paragraph.text)
        if not match:
            continue

        section_num = match.group(1)
        logger.info(f"Found placeholder: {section_num}")

        if section_num in MANUAL_ENTRY_SECTIONS:
            logger.info(
                f"Section {section_num}: Module 1 admin - leave for manual entry."
            )
            continue

        if preserve_template_tables and section_num in OVERLAY_MANAGED_SECTIONS:
            logger.info(
                f"Section {section_num}: handled by QIS v2 overlay - clearing legacy placeholder."
            )
            paragraph.clear()
            processed_sections.add(section_num)
            continue

        if section_num not in section_map:
            logger.warning(
                f"Section {section_num}: no source PDF mapped. Inserting warning."
            )
            _insert_warning(paragraph, section_num)
            warnings_count += 1
            failures += 1
            continue

        if section_num in processed_sections:
            logger.info(
                f"Section {section_num}: duplicate placeholder - clearing."
            )
            paragraph.clear()
            continue

        pdf_path = section_map[section_num]
        logger.info(
            f"Processing {section_num} from '{os.path.basename(pdf_path)}'"
        )

        try:
            content = extract_pdf_content(
                pdf_path            = pdf_path,
                log_folder          = log_folder,
                section_num         = section_num,
                section_page_limits = section_page_limits,
                section_start_pages = section_start_pages,
            )

            current_anchor = paragraph._p
            if section_num == "3.2.S.6":
                # For this section, inject AFTER the static template description line.
                paragraph_index = None
                for idx, p in enumerate(doc.paragraphs):
                    if p._p == paragraph._p:
                        paragraph_index = idx
                        break
                if paragraph_index is not None:
                    for p in doc.paragraphs[paragraph_index + 1 : paragraph_index + 8]:
                        if "description of the container closure system" in p.text.strip().lower():
                            current_anchor = p._p
                            break
            elif section_num == "3.2.P.3.3":
                # Put narrative text under (b). Flow diagram (a) is inserted as image.
                paragraph_index = None
                for idx, p in enumerate(doc.paragraphs):
                    if p._p == paragraph._p:
                        paragraph_index = idx
                        break
                if paragraph_index is not None:
                    for p in doc.paragraphs[paragraph_index + 1 : paragraph_index + 10]:
                        low = p.text.strip().lower()
                        if low.startswith("(b)") and "narrative description" in low:
                            current_anchor = p._p
                            break
            elif section_num == "3.2.P.3.4":
                # Insert extracted controls tables after the template (a) summary line.
                paragraph_index = None
                for idx, p in enumerate(doc.paragraphs):
                    if p._p == paragraph._p:
                        paragraph_index = idx
                        break
                if paragraph_index is not None:
                    for p in doc.paragraphs[paragraph_index + 1 : paragraph_index + 14]:
                        low = p.text.strip().lower()
                        if low.startswith("(a)") and "summary of controls performed at the critical steps" in low:
                            current_anchor = p._p
                            break

            paragraph.clear()

            if content.docx_path and os.path.exists(content.docx_path):
                template_section = _find_template_section_for_paragraph(paragraph)
                if template_section == "2.3.S.4.1":
                    _remove_tables_until_next_section(
                        anchor_xml=current_anchor,
                        logger=logger,
                        section_num=section_num
                    )
                table_keyword = table_keyword_by_template_section.get(
                    template_section, ""
                )
                use_table_only = table_only_all_sections or (
                    section_num in table_only_sections
                )
                if table_keyword:
                    use_table_only = True
                tables_before = len(doc.tables)
                current_anchor = _inject_docx_content(
                    src_docx_path = content.docx_path,
                    anchor_p_xml  = current_anchor,
                    blocklist     = content.noise_blocklist,
                    logger        = logger,
                    section_num   = section_num,
                    include_pdf_tables = include_pdf_tables,
                    table_only    = use_table_only,
                    table_text_keyword = table_keyword
                )
                tables_after = len(doc.tables)
                injected_table_count = tables_after - tables_before

                # Extra Camelot extraction for scanned table near
                # "Lists of Tests and/& Specification" in API spec sections.
                if template_section == "2.3.S.4.1":
                    scanned_rows = _extract_scanned_spec_rows_with_camelot(
                        pdf_path, logger
                    )
                    current_anchor = _append_rows_as_table(
                        doc=doc,
                        current_anchor=current_anchor,
                        headers=["Test", "Acceptance criteria", "Method"],
                        rows=scanned_rows,
                        force_new_table=True,
                        insert_page_break=True
                    )
                    current_anchor = _add_section_spacing(
                        anchor_xml=current_anchor,
                        lines=2
                    )
                try:
                    os.remove(content.docx_path)
                except Exception:
                    pass
            else:
                logger.warning(
                    f"Section {section_num}: no converted DOCX available."
                )

            if section_num == "3.2.P.3.3":
                flow_img_path = _extract_p33_flow_diagram_image(pdf_path, log_folder, logger)
                if flow_img_path and os.path.exists(flow_img_path):
                    inserted = _insert_p33_flow_diagram_image(doc, flow_img_path, logger)
                    if not inserted:
                        logger.warning("Section 3.2.P.3.3: flow-diagram prompt not found in template.")
                    try:
                        os.remove(flow_img_path)
                    except Exception:
                        pass

            sections_filled += 1
            processed_sections.add(section_num)
            logger.info(f"Section {section_num}: populated successfully.")

        except Exception as e:
            logger.error(
                f"Section {section_num} failed: {e}", exc_info=True
            )
            _insert_warning(paragraph, section_num)
            warnings_count += 1
            failures += 1

    logger.info(f"Saving output to {output_path}")
    try:
        # Safe in all modes: fixes converter-introduced zero-width tables.
        _fix_zero_width_tables(doc, logger)

        if not preserve_template_tables:
            _remove_noise_tables(doc, logger)
            _remove_empty_visual_tables(doc, logger)
            _remove_repeated_header_paragraphs(doc, logger)
            _remove_pdf_noise_paragraphs(doc, logger)
            _collapse_blank_paragraphs(doc, logger)
        else:
            logger.info(
                "Template-preserve mode enabled: skipping global cleanup passes "
                "that may remove static QIS tables."
            )
            _remove_low_content_injected_tables(doc, logger, template_table_count)

        logger.info(f"Saving output to {output_path}")
        doc.save(output_path)
    except Exception as e:
        logger.error(f"Failed to save: {e}")
        raise

    return sections_filled, warnings_count, failures
